import unittest
from fpdf import FPDF
from gcloud_cache.cache import cache_result
from docx import Document
import io

class TestCache(unittest.TestCase):

    @cache_result
    def generate_pdf(self, content):
        """Generuje plik PDF i zwraca jego zawartość jako bytes."""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=16)
        
        # Konwertuj binarne dane wejściowe na tekst, jeśli to konieczne
        if isinstance(content, bytes):
            content = content.decode('utf-8')
        
        pdf.cell(200, 10, txt=content, ln=True, align='C')
        
        pdf_path = "/tmp/test_cache.pdf"
        pdf.output(pdf_path)

        with open(pdf_path, "rb") as f:
            pdf_content = f.read()  # Zwracamy zawartość PDF, a nie ścieżkę
        
        return pdf_content

    @cache_result
    def generate_pdf_from_docx(self, doc_bin: bytes):
        """Generuje plik PDF z pliku DOCX i zwraca jego zawartość jako bytes."""
        doc = Document(io.BytesIO(doc_bin))
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        for para in doc.paragraphs:
            pdf.multi_cell(0, 10, para.text)

        pdf_path = "/tmp/test_cache_docx.pdf"
        pdf.output(pdf_path)

        with open(pdf_path, "rb") as f:
            pdf_content = f.read()  # Zwracamy zawartość PDF, a nie ścieżkę
        
        return pdf_content

    def test_pdf_cache(self):
        """Sprawdza, czy plik PDF jest poprawnie cache'owany."""
        pdf1 = self.generate_pdf("Hello, World!")
        pdf2 = self.generate_pdf("Hello, World!")

        self.assertEqual(pdf1, pdf2)  # Porównujemy zawartość, a nie ścieżkę
        self.assertTrue(len(pdf1) > 0)  # Upewniamy się, że PDF nie jest pusty

        print(f"✅ PDF cache test passed. Cached file size: {len(pdf1)} bytes")

    def test_pdf_cache_with_binary_param(self):
        """Sprawdza, czy plik PDF jest poprawnie cache'owany z binarnymi parametrami."""
        binary_content = b"Hello, Binary World!"
        pdf1 = self.generate_pdf(binary_content)
        pdf2 = self.generate_pdf(binary_content)

        self.assertEqual(pdf1, pdf2)  # Porównujemy zawartość, a nie ścieżkę
        self.assertTrue(len(pdf1) > 0)  # Upewniamy się, że PDF nie jest pusty

        print(f"✅ PDF cache with binary param test passed. Cached file size: {len(pdf1)} bytes")

    def test_pdf_from_docx_cache(self):
        """Sprawdza, czy plik PDF generowany z DOCX jest poprawnie cache'owany."""
        doc = Document()
        doc.add_paragraph("Hello, DOCX World!")
        doc_bin = io.BytesIO()
        doc.save(doc_bin)
        doc_bin = doc_bin.getvalue()

        pdf1 = self.generate_pdf_from_docx(doc_bin)
        pdf2 = self.generate_pdf_from_docx(doc_bin)

        self.assertEqual(pdf1, pdf2)  # Porównujemy zawartość, a nie ścieżkę
        self.assertTrue(len(pdf1) > 0)  # Upewniamy się, że PDF nie jest pusty

        print(f"✅ PDF from DOCX cache test passed. Cached file size: {len(pdf1)} bytes")

    def test_cache_from_bin_docx(self):
        """Sprawdza, czy plik PDF generowany z binarnych danych DOCX jest poprawnie cache'owany."""
        doc = Document()
        doc.add_paragraph("Hello, Binary DOCX World!")
        doc_bin = io.BytesIO()
        doc.save(doc_bin)
        doc_bin = doc_bin.getvalue()

        pdf1 = self.generate_pdf_from_docx(doc_bin)
        pdf2 = self.generate_pdf_from_docx(doc_bin)

        self.assertEqual(pdf1, pdf2)  # Porównujemy zawartość, a nie ścieżkę
        self.assertTrue(len(pdf1) > 0)  # Upewniamy się, że PDF nie jest pusty

        print(f"✅ PDF from binary DOCX cache test passed. Cached file size: {len(pdf1)} bytes")

if __name__ == '__main__':
    unittest.main()